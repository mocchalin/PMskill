"""
annotate_prd.py — PRD 红字批注脚本

支持输入格式：
  - .docx   直接处理
  - .md     先转换为 .docx（解析 Markdown 标题/列表/段落）
  - .txt    先转换为 .docx（纯文本逐段落处理）
  - 纯文本  通过 --text 参数传入字符串内容（会先存为 .txt 再转换）

用法：
    python annotate_prd.py <input.docx|.md|.txt> <output.docx> <issues.json>
    python annotate_prd.py --text "需求文本内容..." <output.docx> <issues.json>

issues.json 格式：
{
  "doc_name": "...",
  "conclusion": "需要修改",
  "rating": "⭐⭐",
  "exec_summary": "...",
  "issues": [
    {
      "id": "C-001",
      "dimension": "完整性",
      "severity": "高",           # 高 / 中 / 低
      "anchor_text": "...",       # 用于定位的原文片段（唯一性强）
      "location_desc": "第4章 非功能性需求",
      "problem": "问题描述",
      "suggestion": "修改建议"
    },
    ...
  ]
}
"""

import sys
import re
import os
import json
import datetime
import tempfile
from copy import deepcopy
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Input conversion helpers ──────────────────────────────────────────────────

def markdown_to_docx(md_text: str, output_path: str):
    """
    Convert Markdown text to a .docx file.
    Handles: # headings, **bold**, *italic*, - / * / 1. lists, plain paragraphs,
             --- horizontal rules (as paragraph breaks), and code blocks (as monospace).
    """
    doc = Document()

    # Trim leading/trailing whitespace per line, split into lines
    lines = md_text.splitlines()
    in_code_block = False
    code_lines = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # Code block fence (``` or ~~~)
        if re.match(r'^(`{3}|~{3})', line):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                # Add code block as monospace paragraphs
                for cl in code_lines:
                    p = doc.add_paragraph(style='Normal')
                    run = p.add_run(cl)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    p.paragraph_format.left_indent = Cm(0.8)
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,6})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            # Map to Word heading styles (Heading 1-3; deeper levels use Heading 3)
            style_name = f'Heading {min(level, 3)}'
            try:
                doc.add_heading(text, level=min(level, 3))
            except Exception:
                doc.add_paragraph(text, style='Normal').runs[0].font.bold = True
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^[-*_]{3,}\s*$', line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # Unordered list items (- or * or +)
        ul_match = re.match(r'^[\-\*\+]\s+(.*)', line)
        if ul_match:
            item_text = ul_match.group(1)
            p = doc.add_paragraph(style='List Bullet')
            _add_inline_formatting(p, item_text)
            i += 1
            continue

        # Ordered list items
        ol_match = re.match(r'^\d+\.\s+(.*)', line)
        if ol_match:
            item_text = ol_match.group(1)
            p = doc.add_paragraph(style='List Number')
            _add_inline_formatting(p, item_text)
            i += 1
            continue

        # Blank line → paragraph break (skip)
        if line.strip() == '':
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph(style='Normal')
        _add_inline_formatting(p, line)
        i += 1

    doc.save(output_path)
    return output_path


def _add_inline_formatting(para, text: str):
    """
    Parse inline **bold**, *italic*, `code` and add as styled runs.
    Falls back to plain text for unrecognized patterns.
    """
    # Pattern matches **bold**, *italic*, `code`, and plain text segments
    token_re = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|([^*`]+))')
    for m in token_re.finditer(text):
        if m.group(2):  # **bold**
            run = para.add_run(m.group(2))
            run.font.bold = True
        elif m.group(3):  # *italic*
            run = para.add_run(m.group(3))
            run.font.italic = True
        elif m.group(4):  # `code`
            run = para.add_run(m.group(4))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        elif m.group(5):  # plain text
            para.add_run(m.group(5))


def plaintext_to_docx(text: str, output_path: str):
    """
    Convert plain text to .docx. Each non-empty line group becomes a paragraph.
    Lines that look like headings (ALL CAPS, or ending with colon, or numbered like "1.")
    are styled as bold. Indented lines are treated as list items.
    """
    doc = Document()
    paragraphs = re.split(r'\n{2,}', text.strip())  # split on blank lines

    for block in paragraphs:
        lines = block.strip().splitlines()
        if not lines:
            continue

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            # Detect heading-like lines: starts with numbers/letters + dot, or ALL CAPS short line
            if re.match(r'^\d+[\.\)]\s', stripped) and len(stripped) < 80:
                p = doc.add_paragraph(style='Normal')
                run = p.add_run(stripped)
                run.font.bold = True
            elif stripped.isupper() and len(stripped) < 60:
                p = doc.add_paragraph(style='Heading 2')
                p.add_run(stripped)
            elif line.startswith('    ') or line.startswith('\t'):
                # Indented → list bullet
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(stripped)
            else:
                doc.add_paragraph(stripped, style='Normal')

    doc.save(output_path)
    return output_path


def ensure_docx(input_path: str) -> tuple:
    """
    Given any supported input path, return (docx_path, is_temp).
    If input is already .docx, return as-is (is_temp=False).
    Otherwise convert and return path to a temp .docx (is_temp=True).
    """
    ext = os.path.splitext(input_path)[1].lower()

    if ext == '.docx':
        return input_path, False

    # Read the source text
    with open(input_path, encoding='utf-8', errors='replace') as f:
        content = f.read()

    # Write to a temp .docx
    tmp_fd, tmp_path = tempfile.mkstemp(suffix='.docx')
    os.close(tmp_fd)

    if ext == '.md':
        markdown_to_docx(content, tmp_path)
    else:
        # .txt or any other text format
        plaintext_to_docx(content, tmp_path)

    return tmp_path, True


def text_to_docx_temp(text: str, is_markdown: bool = False) -> str:
    """
    Convert a raw text string to a temporary .docx file.
    Returns the path to the temp file.
    """
    tmp_fd, tmp_path = tempfile.mkstemp(suffix='.docx')
    os.close(tmp_fd)

    if is_markdown:
        markdown_to_docx(text, tmp_path)
    else:
        plaintext_to_docx(text, tmp_path)

    return tmp_path


# ── Emoji / Unicode safety ─────────────────────────────────────────────────

def _detect_emoji_support() -> bool:
    """Detect if the runtime supports 4-byte UTF-8 emoji (U+1F300+).

    Some inference backends or terminal environments truncate or mangle
    characters outside the BMP. When that happens the docx XML ends up
    with replacement characters (�) which display as tofu / boxes.
    """
    try:
        test = "\U0001f534\U0001f7e1\U0001f7e2"  # 🔴🟡🟢
        return test.encode("utf-8").decode("utf-8") == test
    except (UnicodeEncodeError, UnicodeDecodeError):
        return False


_EMOJI_OK = _detect_emoji_support()

# Safe fall-backs for every special glyph used in the script
_SAFE = {
    "summary_icon": "\U0001f4cb" if _EMOJI_OK else "[汇总]",       # 📋
    "arrow":        "\u2192"    if _EMOJI_OK else "->",            # →
    "ellipsis":     "\u2026"    if _EMOJI_OK else "...",           # …
    "bar":          "\u258c"    if _EMOJI_OK else "||",            # ▌
    "bullet":       "\u25b6"    if _EMOJI_OK else ">>",            # ▶
    "problem":      "\u2757"    if _EMOJI_OK else "[!]",           # ❗
    "suggest":      "\u2705"    if _EMOJI_OK else "[OK]",          # ✅
    "done":         "\u2705"    if _EMOJI_OK else "[OK]",          # ✅
    "star":         "\u2b50"    if _EMOJI_OK else "*",             # ⭐
}

# ── Color constants ──────────────────────────────────────────────────────────
RED       = RGBColor(0xCC, 0x00, 0x00)   # standard red annotation
ORANGE    = RGBColor(0xC0, 0x50, 0x00)   # medium priority
DARK_RED  = RGBColor(0x99, 0x00, 0x00)   # low priority (still red family)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
NAVY      = RGBColor(0x1F, 0x35, 0x64)

SEV_META = {
    "高": {"icon": "\U0001f534" if _EMOJI_OK else "[高]", "color": RED,      "bg": "FFE8E8", "label": "高优先级"},
    "中": {"icon": "\U0001f7e1" if _EMOJI_OK else "[中]", "color": ORANGE,   "bg": "FFF4E0", "label": "中优先级"},
    "低": {"icon": "\U0001f7e2" if _EMOJI_OK else "[低]", "color": DARK_RED, "bg": "F0FFE8", "label": "低优先级"},
}

DIM_CODE = {
    "完整性": "C",
    "一致性": "CO",
    "清晰性": "CL",
    "可测试性": "T",
}


# ── XML helpers ──────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_left_border(para, hex_color: str = "CC0000", width_eighths: int = 36):
    """Add a left border (red stripe) to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(width_eighths))
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), hex_color)
    pBdr.append(left)
    pPr.append(pBdr)


def add_para_bg(para, hex_color: str):
    """Set paragraph shading background."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def add_bookmark(para, bm_id: int, bm_name: str):
    """Wrap the paragraph content with a bookmark."""
    p_elem = para._p
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), str(bm_id))
    bm_start.set(qn('w:name'), bm_name)
    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), str(bm_id))
    # Insert at the very beginning of the paragraph
    p_elem.insert(0, bm_start)
    p_elem.append(bm_end)


def add_internal_hyperlink(para, anchor: str, display_text: str, color: RGBColor = None):
    """Append an internal hyperlink (jumps to bookmark) to a paragraph."""
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), anchor)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # Color
    c = OxmlElement('w:color')
    if color:
        c.set(qn('w:val'), '{:02X}{:02X}{:02X}'.format(*color))
    else:
        c.set(qn('w:val'), '0563C1')
    rPr.append(c)

    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = display_text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    hyperlink.append(r)
    para._p.append(hyperlink)
    return hyperlink


def insert_paragraph_after(ref_paragraph):
    """Insert a blank paragraph immediately after ref_paragraph and return it."""
    new_p = OxmlElement('w:p')
    ref_paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, ref_paragraph._p.getparent())


def insert_page_break_para_before_first(body_elem):
    """Insert a page-break paragraph before the first body element."""
    pb_p = OxmlElement('w:p')
    pb_r = OxmlElement('w:r')
    pb_br = OxmlElement('w:br')
    pb_br.set(qn('w:type'), 'page')
    pb_r.append(pb_br)
    pb_p.append(pb_r)
    # Insert just before position 0 (will become index 1 after summary)
    return pb_p


# ── Annotation helpers ───────────────────────────────────────────────────────

def make_annotation_text(issue: dict) -> str:
    sev = issue["severity"]
    icon = SEV_META[sev]["icon"]
    label = SEV_META[sev]["label"]
    return (
        f"{_SAFE['bar']} 审核批注 [{issue['id']} | {icon} {label} | {issue['dimension']}]\n"
        f"{_SAFE['problem']} 问题：{issue['problem']}\n"
        f"{_SAFE['suggest']} 建议：{issue['suggestion']}"
    )


def add_annotation_paragraph(ref_para, issue: dict, bm_id: int) -> bool:
    """
    Insert a red-text annotation paragraph after ref_para.
    Returns True on success.
    """
    sev = issue["severity"]
    meta = SEV_META[sev]
    color = meta["color"]
    bg = meta["bg"]
    bm_name = "issue_" + issue["id"].replace("-", "_")

    anno_para = insert_paragraph_after(ref_para)

    # Background shading
    add_para_bg(anno_para, bg)
    # Left red border
    add_left_border(anno_para, "CC0000", 36)

    # Paragraph spacing
    pf = anno_para.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.left_indent = Cm(0.3)

    # Header line
    header_run = anno_para.add_run(
        f"{_SAFE['bar']} 审核批注 [{issue['id']} | {meta['icon']} {meta['label']} | {issue['dimension']}]"
    )
    header_run.font.color.rgb = color
    header_run.font.bold = True
    header_run.font.italic = True
    header_run.font.size = Pt(9)

    # Problem line
    anno_para.add_run("\n")
    prob_label = anno_para.add_run(f"{_SAFE['problem']} 问题：")
    prob_label.font.color.rgb = color
    prob_label.font.bold = True
    prob_label.font.italic = True
    prob_label.font.size = Pt(9)
    prob_text = anno_para.add_run(issue["problem"])
    prob_text.font.color.rgb = color
    prob_text.font.italic = True
    prob_text.font.size = Pt(9)

    # Suggestion line
    anno_para.add_run("\n")
    sugg_label = anno_para.add_run(f"{_SAFE['suggest']} 建议：")
    sugg_label.font.color.rgb = color
    sugg_label.font.bold = True
    sugg_label.font.italic = True
    sugg_label.font.size = Pt(9)
    sugg_text = anno_para.add_run(issue["suggestion"])
    sugg_text.font.color.rgb = color
    sugg_text.font.italic = True
    sugg_text.font.size = Pt(9)

    # Add bookmark to this annotation paragraph
    add_bookmark(anno_para, bm_id, bm_name)

    return True


# ── Paragraph search ─────────────────────────────────────────────────────────

def get_all_paragraphs(doc):
    """
    Return all paragraphs in document order, including those inside tables.
    Each item: (paragraph_object, is_in_table)
    """
    result = []
    body = doc.element.body

    def _collect(element):
        from docx.oxml.ns import qn
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        if tag == 'p':
            from docx.text.paragraph import Paragraph
            result.append(Paragraph(element, element.getparent()))
        else:
            for child in element:
                _collect(child)

    _collect(body)
    return result


def find_anchor_paragraph(paragraphs, anchor_text: str):
    """
    Find the paragraph that contains anchor_text.
    Returns the paragraph object or None.
    Strategy: try exact substring match first, then fuzzy (first 15 chars).
    """
    anchor_clean = anchor_text.strip()

    # 1. Full substring match
    for para in paragraphs:
        if anchor_clean in para.text:
            return para

    # 2. Partial match with first 20 chars of anchor (more tolerant)
    short = anchor_clean[:20]
    if len(short) >= 8:
        for para in paragraphs:
            if short in para.text:
                return para

    # 3. Word-by-word: all significant words (>=3 chars) present
    words = [w for w in anchor_clean.split() if len(w) >= 3]
    if words:
        for para in paragraphs:
            if all(w in para.text for w in words):
                return para

    return None


# ── Summary page builder ─────────────────────────────────────────────────────

def build_summary_section(doc, issues: list, exec_summary: str, conclusion: str, rating: str, doc_name: str):
    """
    Prepend a summary index page to the document (inserted before everything else).
    Returns the list of elements inserted (for ordering).
    """
    body = doc.element.body

    # We'll collect all new XML elements and then insert them at the top in reverse order
    # (since each insert(0, ...) prepends)

    # ── Page break at the end of summary (separates from original content) ──
    pb_p = insert_page_break_para_before_first(body)
    body.insert(0, pb_p)   # will be pushed down as we prepend more

    # We need to build the summary elements and insert them in order.
    # Easiest: use a temporary Document, build content there, then transplant XML nodes.
    from docx import Document as DocClass
    tmp = DocClass()
    # Remove the default empty paragraph
    for p in tmp.paragraphs:
        p._element.getparent().remove(p._element)

    # ── Title ──
    title_p = tmp.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run(f"{_SAFE['summary_icon']}  PRD 审核问题汇总")
    tr.font.size = Pt(18)
    tr.font.bold = True
    tr.font.color.rgb = NAVY

    # ── Meta info ──
    meta_p = tmp.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta_p.add_run(
        f"文档：{doc_name}   |   审核日期：{datetime.date.today().strftime('%Y-%m-%d')}"
        f"   |   结论：{conclusion}  {rating}"
    )
    mr.font.size = Pt(10)
    mr.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    tmp.add_paragraph()  # spacer

    # ── Executive summary ──
    es_title = tmp.add_paragraph()
    estr = es_title.add_run(f"{_SAFE['bullet']} 执行摘要")
    estr.font.bold = True
    estr.font.size = Pt(11)
    estr.font.color.rgb = NAVY

    es_body = tmp.add_paragraph(exec_summary)
    es_body.paragraph_format.left_indent = Cm(0.5)
    for run in es_body.runs:
        run.font.size = Pt(10)

    tmp.add_paragraph()  # spacer

    # ── Issue count summary ──
    cnt_high = sum(1 for i in issues if i["severity"] == "高")
    cnt_mid  = sum(1 for i in issues if i["severity"] == "中")
    cnt_low  = sum(1 for i in issues if i["severity"] == "低")

    cnt_p = tmp.add_paragraph()
    hi = SEV_META["高"]["icon"]
    mi = SEV_META["中"]["icon"]
    lo = SEV_META["低"]["icon"]
    cnt_r = cnt_p.add_run(
        f"共发现问题 {len(issues)} 项：{hi} 高优先级 {cnt_high} 项  |  {mi} 中优先级 {cnt_mid} 项  |  {lo} 低优先级 {cnt_low} 项"
    )
    cnt_r.font.size = Pt(10)
    cnt_r.font.bold = True

    tmp.add_paragraph()

    # ── Index table ──
    idx_title = tmp.add_paragraph()
    itr = idx_title.add_run(f"{_SAFE['bullet']} 问题索引（点击'跳转'可直达批注位置）")
    itr.font.bold = True
    itr.font.size = Pt(11)
    itr.font.color.rgb = NAVY

    tbl = tmp.add_table(rows=1, cols=5)
    tbl.style = 'Table Grid'

    # Header row
    headers = ["编号", "严重程度", "维度", "问题简述", "位置 / 跳转"]
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, "1F3564")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.color.rgb = WHITE
        run.font.bold = True
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Set column widths
    col_widths = [Cm(1.5), Cm(2.2), Cm(2.2), Cm(9.0), Cm(3.5)]
    for i, w in enumerate(col_widths):
        for cell in tbl.columns[i].cells:
            cell.width = w

    # Data rows
    for issue in issues:
        sev = issue["severity"]
        meta = SEV_META[sev]
        row_cells = tbl.add_row().cells

        # Col 0: ID
        row_cells[0].paragraphs[0].add_run(issue["id"]).font.size = Pt(9)
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Col 1: Severity
        sev_p = row_cells[1].paragraphs[0]
        sev_r = sev_p.add_run(f"{meta['icon']} {meta['label']}")
        sev_r.font.color.rgb = meta["color"]
        sev_r.font.bold = True
        sev_r.font.size = Pt(9)
        set_cell_bg(row_cells[1], meta["bg"])

        # Col 2: Dimension
        row_cells[2].paragraphs[0].add_run(issue["dimension"]).font.size = Pt(9)

        # Col 3: Problem summary (truncated)
        problem_short = issue["problem"][:60] + (_SAFE["ellipsis"] if len(issue["problem"]) > 60 else "")
        row_cells[3].paragraphs[0].add_run(problem_short).font.size = Pt(9)

        # Col 4: Location + jump link
        loc_p = row_cells[4].paragraphs[0]
        loc_p.add_run(issue["location_desc"] + "  ").font.size = Pt(9)
        bm_name = "issue_" + issue["id"].replace("-", "_")
        add_internal_hyperlink(loc_p, bm_name, f"{_SAFE['arrow']} 跳转")

    # ── Transplant all tmp elements into main doc body at position 0 ──
    tmp_body = tmp.element.body
    # Collect all child elements of tmp body (except the final sectPr if present)
    tmp_children = [child for child in tmp_body if child.tag != qn('w:sectPr')]

    # Insert in reverse order at position 0 to maintain order
    for elem in reversed(tmp_children):
        body.insert(0, deepcopy(elem))

    # The page break (pb_p) is now at some position; we need to ensure it's
    # right after the summary content and before original content.
    # Find pb_p's current position and move it right after last summary element.
    # Since we inserted pb_p at 0 first, then prepended summary elements,
    # the page break is now at index len(tmp_children).
    # That's correct — summary elements are 0..N-1, page break is at N, original content follows.


# ── Main ─────────────────────────────────────────────────────────────────────

def annotate(input_path: str, output_path: str, issues: list,
             exec_summary: str = "", conclusion: str = "需要修改",
             rating: str = "⭐⭐", doc_name: str = "",
             raw_text: str = None, is_markdown: bool = False):
    """
    Main annotation entry point.

    Parameters
    ----------
    input_path  : path to source document (.docx / .md / .txt).
                  Pass None if providing raw_text instead.
    raw_text    : raw string content (used when user pastes text directly).
                  If provided, input_path is ignored.
    is_markdown : set True when raw_text contains Markdown syntax.
    """
    tmp_docx = None  # track temp file for cleanup

    try:
        # ── Step 1: resolve input to a .docx ──
        if raw_text is not None:
            tmp_docx = text_to_docx_temp(raw_text, is_markdown=is_markdown)
            docx_path = tmp_docx
            if not doc_name:
                doc_name = "粘贴文本"
        else:
            docx_path, is_temp = ensure_docx(input_path)
            if is_temp:
                tmp_docx = docx_path
            if not doc_name:
                doc_name = os.path.basename(input_path)

        doc = Document(docx_path)
        all_paras = get_all_paragraphs(doc)

        # ── Step 2: locate anchors ──
        placed = []
        unplaced = []
        bm_counter = 100

        located = []
        for issue in issues:
            para = find_anchor_paragraph(all_paras, issue["anchor_text"])
            if para:
                located.append((issue, para))
            else:
                unplaced.append(issue)

        # Sort by document order (reverse) so bottom-up insertion keeps positions stable
        body = doc.element.body
        all_p_elements = list(body.iter(qn('w:p')))

        def para_index(para):
            try:
                return all_p_elements.index(para._p)
            except ValueError:
                return 999999

        located.sort(key=lambda x: para_index(x[1]), reverse=True)

        # ── Step 3: insert annotations bottom-up ──
        for issue, ref_para in located:
            add_annotation_paragraph(ref_para, issue, bm_counter)
            bm_counter += 1
            placed.append(issue)

        # ── Step 4: handle unplaced issues at end of doc ──
        if unplaced:
            end_title = doc.add_paragraph()
            end_title.add_run("── 无法定位到原文的批注 ──").font.color.rgb = RED
            for issue in unplaced:
                ep = doc.add_paragraph()
                add_para_bg(ep, SEV_META[issue["severity"]]["bg"])
                add_left_border(ep, "CC0000", 36)
                r = ep.add_run(
                    f"[{issue['id']} | {SEV_META[issue['severity']]['icon']} {issue['severity']} | {issue['dimension']}]  "
                    f"{_SAFE['problem']} {issue['problem']}  {_SAFE['suggest']} {issue['suggestion']}"
                )
                r.font.color.rgb = RED
                r.font.italic = True
                r.font.size = Pt(9)
                add_bookmark(ep, bm_counter, "issue_" + issue["id"].replace("-", "_"))
                bm_counter += 1

        # ── Step 5: prepend summary index page ──
        build_summary_section(doc, issues, exec_summary, conclusion, rating, doc_name)

        doc.save(output_path)
        return {"placed": len(placed), "unplaced": len(unplaced), "total": len(issues)}

    finally:
        # Clean up any temporary .docx we created from .md/.txt/raw text
        if tmp_docx and os.path.exists(tmp_docx):
            try:
                os.remove(tmp_docx)
            except Exception:
                pass


# ── CLI entry point ───────────────────────────────────────────────────────────
if __name__ == "__main__":
    # Support: python annotate_prd.py --text "..." output.docx issues.json [--markdown]
    # Support: python annotate_prd.py input.docx output.docx issues.json
    args = sys.argv[1:]

    if len(args) >= 3 and args[0] == '--text':
        raw_text    = args[1]
        output_path = args[2]
        issues_path = args[3] if len(args) > 3 else None
        is_markdown = '--markdown' in args
        input_path  = None
    elif len(args) >= 3:
        input_path  = args[0]
        output_path = args[1]
        issues_path = args[2]
        raw_text    = None
        is_markdown = False
    else:
        print("Usage:")
        print("  python annotate_prd.py <input.docx|.md|.txt> <output.docx> <issues.json>")
        print("  python annotate_prd.py --text 'PRD content...' <output.docx> <issues.json> [--markdown]")
        sys.exit(1)

    with open(issues_path, encoding="utf-8") as f:
        data = json.load(f)

    issues       = data["issues"]
    exec_summary = data.get("exec_summary", "")
    conclusion   = data.get("conclusion", "需要修改")
    rating       = data.get("rating", "⭐⭐")
    doc_name     = data.get("doc_name", "")

    result = annotate(
        input_path=input_path,
        output_path=output_path,
        issues=issues,
        exec_summary=exec_summary,
        conclusion=conclusion,
        rating=rating,
        doc_name=doc_name,
        raw_text=raw_text,
        is_markdown=is_markdown,
    )
    print(f"{_SAFE['done']} 完成：成功定位 {result['placed']} 条，未定位 {result['unplaced']} 条，共 {result['total']} 条")
    print(f"   输出文件：{output_path}")
